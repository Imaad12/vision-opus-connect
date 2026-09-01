"""WordImporter tests, using a small synthetic .docx built with
python-docx at test time."""

from __future__ import annotations

from pathlib import Path

import docx

from app.importers.word_importer import WordImporter


def test_docx_import_extracts_headings_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "quote.docx"
    document = docx.Document()
    document.add_heading("Quotation Q-2024-0091", level=1)
    document.add_paragraph("Please find our quotation for the referenced project.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Description"
    table.rows[0].cells[1].text = "Amount"
    table.rows[1].cells[0].text = "Excavation works"
    table.rows[1].cells[1].text = "5000.00"
    document.save(path)

    result = WordImporter().extract(path)

    assert not result.unsupported
    assert "Quotation Q-2024-0091" in result.text
    assert "Please find our quotation" in result.text
    assert len(result.tables) == 1
    assert result.tables[0].rows[1] == ["Excavation works", "5000.00"]


def test_legacy_doc_is_reported_as_unsupported_not_faked(tmp_path: Path) -> None:
    path = tmp_path / "old_quote.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0 pretend legacy OLE content")

    result = WordImporter().extract(path)

    assert result.unsupported is True
    assert ".doc" in (result.unsupported_reason or "")


def test_corrupt_docx_is_reported_as_unsupported_not_crashed(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a real docx file")

    result = WordImporter().extract(path)

    assert result.unsupported is True
    assert result.unsupported_reason


def test_empty_docx_reports_a_warning(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    docx.Document().save(path)

    result = WordImporter().extract(path)

    assert not result.unsupported
    assert result.text is None
    assert any("empty" in warning.lower() for warning in result.warnings)
