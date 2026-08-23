"""Word importer: .docx via python-docx. Legacy .doc is explicitly not
supported — it is a different (binary OLE) file format that python-docx
cannot open, and Phase 4 does not bundle a converter. Per the brief ("do
not fake support"), a .doc file is reported as needing conversion, never
silently misread as empty or crashed on.
"""

from __future__ import annotations

from pathlib import Path

import docx

from app.importers.base import BaseImporter, ExtractedTable, RawExtraction


class WordImporter(BaseImporter):
    extensions = ("docx", "doc")

    def extract(self, path: Path) -> RawExtraction:
        if path.suffix.lower() == ".doc":
            return RawExtraction(
                unsupported=True,
                unsupported_reason=(
                    "Legacy .doc files are not supported in Phase 4 — this is a different "
                    "binary format from .docx and needs a separate converter. Save this file "
                    "as .docx (e.g. from Word or LibreOffice) and re-import."
                ),
            )

        try:
            document = docx.Document(str(path))
        except Exception as exc:  # noqa: BLE001
            return RawExtraction(unsupported=True, unsupported_reason=f"Could not open .docx file: {exc}")

        text_parts: list[str] = []
        for paragraph in document.paragraphs:
            content = paragraph.text.strip()
            if not content:
                continue
            if paragraph.style is not None and paragraph.style.name.lower().startswith("heading"):
                text_parts.append(f"# {content}")
            else:
                text_parts.append(content)

        tables: list[ExtractedTable] = []
        for table_index, table in enumerate(document.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            rows = [row for row in rows if any(cell for cell in row)]
            if rows:
                tables.append(ExtractedTable(name=f"table {table_index + 1}", rows=rows))

        warnings: list[str] = []
        if not text_parts and not tables:
            warnings.append("This document appears to be empty.")

        return RawExtraction(text="\n".join(text_parts) or None, tables=tables, warnings=warnings)
